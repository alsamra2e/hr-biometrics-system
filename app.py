import streamlit as st
import pandas as pd
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import requests
from io import BytesIO

# --- HELPER FUNCTIONS ---
def set_rtl(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    bidi = qn('w:bidi')
    if not pPr.xpath(f'./{bidi}'):
        pPr.append(p.make_element(bidi))

def extract_date_from_filename(filename):
    # Regex to find YYYY-MM-DD in the filename
    match = re.search(r'\d{4}-\d{2}-\d{2}', filename)
    return match.group(0) if match else "Unknown Date"

# --- THE NEW EXCEPTIONS MODULE ---
def run_exceptions_module():
    st.subheader("📋 Exceptions & Attendance Audit")
    uploaded_files = st.file_uploader("Upload Daily Excels (Multiple)", accept_multiple_files=True, type=['xlsx', 'xls'])
    
    if uploaded_files:
        all_data = []
        for f in uploaded_files:
            file_date = extract_date_from_filename(f.name)
            df = pd.read_excel(f)
            df['Report_Date'] = file_date # Add date from filename to every row
            all_data.append(df)
        
        combined_df = pd.concat(all_data, ignore_index=True)
        
        # Filter: Keep only Late and Absence (ignoring emojis)
        mask = combined_df['Status'].str.contains('Late|Absence', na=False)
        exceptions = combined_df[mask].copy()
        
        if not exceptions.empty:
            # Grouping by Name and Status to collect dates
            summary = exceptions.groupby(['Name', 'Status'])['Report_Date'].agg(list).reset_index()
            summary['Count'] = summary['Report_Date'].apply(len)
            summary['Dates_Str'] = summary['Report_Date'].apply(lambda x: ", ".join(x))

            st.write(f"Found {len(summary)} records with lates or absences.")
            
            if st.button("Generate Official Alturath Report"):
                with st.spinner("Creating Word Document..."):
                    report_bytes = create_word_doc(summary)
                    st.download_button("📥 Download Official Report", report_bytes, "Alturath_Exceptions_Report.docx")
        else:
            st.success("Perfect! No Lates or Absences found.")

def create_word_doc(df):
    doc = Document()
    
    # 1. HEADER (3 Columns)
    section = doc.sections[0]
    header = section.header
    htable = header.add_table(1, 3, width=Inches(6.5))
    
    # Right: Arabic
    r = htable.rows[0].cells[0].paragraphs[0]
    r.text = "جامعة التراث\nقسم الشؤون الإدارية والمالية\nشعبة الموارد البشرية"
    r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(r)
    
    # Middle: Logo
    m = htable.rows[0].cells[1].paragraphs[0]
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        img_url = "https://uoturath.edu.iq/wp-content/uploads/2025/03/shield-1.png"
        img_data = BytesIO(requests.get(img_url).content)
        m.add_run().add_picture(img_data, width=Inches(0.8))
    except: pass
    
    # Left: English
    l = htable.rows[0].cells[2].paragraphs[0]
    l.text = "University Of Alturath\nDept. Of Admin & Financial Affairs\nHR Department"
    l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 2. SEPARATOR LINE
    p_line = doc.add_paragraph()
    run = p_line.add_run("______________________________________________________________________")
    run.font.color.rgb = RGBColor(0x8F, 0x0B, 0x0B) # Maroon
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. BODY TEXT
    body = doc.add_paragraph("\nنرفق لسيادتكم في ادناه الكشف الخاص بموقف الحضور والغياب لكادر العمل الخاص بجامعة التراث وحسب كشف البصمة المرفق طيا نسخة منه ... راجين التفضل بالاطلاع واعلامنا توجيهات سيادتكم حول ذلك ... مع التقدير..")
    body.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(body)

    # 4. DATA TABLE
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    labels = ['الاسم', 'الحالة', 'العدد', 'التواريخ']
    for i, txt in enumerate(labels):
        hdr[i].text = txt
        set_rtl(hdr[i].paragraphs[0])

    for _, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['Name'])
        cells[1].text = "تأخير" if "Late" in row['Status'] else "غياب"
        cells[2].text = str(row['Count'])
        
        # Small font for dates to keep table clean
        d_para = cells[3].paragraphs[0]
        d_run = d_para.add_run(row['Dates_Str'])
        d_run.font.size = Pt(8)
        set_rtl(d_para)

    # 5. SIGNATURE
    doc.add_paragraph("\n\n")
    sig = doc.add_paragraph("م.م محمد زهير طالب النقيب\nمدير قسم الشؤون الادارية والموارد البشرية")
    sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_rtl(sig)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# --- MAIN APP NAVIGATION ---
st.sidebar.title("HR Management System")
mode = st.sidebar.radio("Go to:", ["Daily Report (Old)", "Exceptions Audit (New)"])

if mode == "Daily Report (Old)":
    st.title("Biometric Attendance Report")
    # PASTE YOUR ENTIRE ORIGINAL CODE HERE
    import streamlit as st
import pandas as pd
import plotly.express as px
from io import BytesIO
from datetime import datetime, date

# 1. PAGE SETUP
st.set_page_config(page_title="Alturath University | HR Audit Pro", layout="wide")

# 2. SIDEBAR LOGO & FILTERS
st.sidebar.image("https://encrypted-tbn0.gstatic.com/images?q=tbn:ANd9GcRfTMmtmrsxGUBnlEb0xB0ClMbFZmj_L5Ap5Q&s")
st.sidebar.title("HR Audit Control")

with st.sidebar:
    st.divider()
    st.subheader("📅 Audit Parameters")
    use_today = st.toggle("Show Today Only", value=False)
    target_date = date.today() if use_today else st.date_input("Audit Date", value=date.today())
    
    weekdays_ar = {
        "Monday": "الاثنين", "Tuesday": "الثلاثاء", "Wednesday": "الاربعاء", 
        "Thursday": "الخميس", "Friday": "الجمعة", "Saturday": "السبت", "Sunday": "الاحد"
    }
    current_weekday_ar = weekdays_ar[target_date.strftime("%A")]
    st.info(f"Audit Day: **{current_weekday_ar}**")

    st.subheader("📥 Data Sources")
    f_zaqura = st.file_uploader("Zaqura Gate", type=['xlsx', 'xls'])
    f_mhmd = st.file_uploader("Mhmd Bn Ali Gate", type=['xlsx', 'xls'])
    f_app = st.file_uploader("Mawjood App", type=['xlsx', 'xls'])
    f_weekly = st.file_uploader("📅 Weekly Day-Off List", type=['xlsx', 'xls'])

# 3. DATA PROCESSING
def process_gate(file, g_name):
    try:
        engine = 'xlrd' if file.name.endswith('.xls') else 'openpyxl'
        df = pd.read_excel(file, engine=engine)
        df.columns = [str(c).strip() for c in df.columns]
        df['dt'] = pd.to_datetime(df['الوقت'], errors='coerce')
        df = df[df['dt'].dt.date == target_date]
        df['Time'] = df['dt'].dt.strftime('%H:%M')
        df = df.rename(columns={'الاسم': 'Name', 'الإسم': 'Name', 'رقم هوية': 'ID'})
        return df[['Name', 'Time']].assign(Source=g_name)
    except: return pd.DataFrame()

def process_app(file):
    try:
        engine = 'xlrd' if file.name.endswith('.xls') else 'openpyxl'
        df = pd.read_excel(file, header=3, engine=engine)
        df.columns = [str(c).strip() for c in df.columns]
        # Important: Match the exact status name for presence
        df = df[df['الحالة'].isin(['حاضر', 'Present'])]
        df['Time'] = pd.to_datetime(df['دخول'], errors='coerce').dt.strftime('%H:%M')
        return pd.DataFrame({'Name': df['الاسم'], 'Time': df['Time'], 'Source': 'App'})
    except: return pd.DataFrame()

def process_weekly_off(file):
    try:
        engine = 'xlrd' if file.name.endswith('.xls') else 'openpyxl'
        df = pd.read_excel(file, engine=engine)
        df.columns = [str(c).strip() for c in df.columns]
        return df.rename(columns={'الاسم الثلاثي': 'Name', 'الاجازة الاسبوعية': 'OffDay'})
    except: return pd.DataFrame()

# 4. CONSOLIDATION & ANALYSIS
all_logs = []
if f_zaqura: all_logs.append(process_gate(f_zaqura, "Zaqura Gate"))
if f_mhmd: all_logs.append(process_gate(f_mhmd, "Mhmd Bn Ali Gate"))
if f_app: all_logs.append(process_app(f_app))

if all_logs or f_weekly:
    # Standardize Logs
    if all_logs:
        df_present = pd.concat(all_logs, ignore_index=True)
        df_present = df_present.dropna(subset=['Time', 'Name'])
        # Keep earliest punch of the day per name
        df_present = df_present.sort_values('Time').drop_duplicates(subset=['Name'], keep='first')
    else:
        df_present = pd.DataFrame(columns=['Name', 'Time', 'Source'])
    
    df_off = process_weekly_off(f_weekly) if f_weekly else pd.DataFrame(columns=['Name', 'OffDay'])
    
    # Unified list of all unique names across logs and staff list
    master_names = list(set(df_present['Name'].tolist() + df_off['Name'].tolist()))
    
    final_data = []
    for name in master_names:
        punch = df_present[df_present['Name'] == name]
        off_info = df_off[df_off['Name'] == name]
        is_off_today = (off_info['OffDay'].iloc[0] == current_weekday_ar) if not off_info.empty else False
        
        row = {"Name": name, "Check-In": "-", "Source": "-", "Status": ""}
        
        # --- LOGIC PRIORITY ---
        if not punch.empty:
            # If they punched in, they are PRESENT (regardless of day-off status)
            row["Check-In"] = punch['Time'].iloc[0]
            row["Source"] = punch['Source'].iloc[0]
            row["Status"] = "🔴 Late" if row["Check-In"] > "08:35" else "✅ On Time"
        elif is_off_today:
            # If they didn't punch but it's their scheduled day off
            row["Status"] = "🟡 Weekly Off"
        else:
            # If they didn't punch and it's a working day
            row["Status"] = "❌ Absence"
            
        final_data.append(row)

    df_final = pd.DataFrame(final_data)

    # 5. UI TABS
    tab1, tab2 = st.tabs(["📊 Performance Analysis", "🕵️ Detailed Presence Log"])

    with tab1:
        st.header(f"Migration & Lateness Progress ({current_weekday_ar})")
        if not df_final.empty:
            m1, m2, m3 = st.columns(3)
            total = len(df_final)
            late = len(df_final[df_final['Status'] == "🔴 Late"])
            absent = len(df_final[df_final['Status'] == "❌ Absence"])
            m1.metric("Staff Coverage", f"{((total-absent)/total)*100:.1f}%" if total > 0 else "0%")
            m2.metric("Lateness Rate", f"{(late/total)*100:.1f}%" if total > 0 else "0%", delta_color="inverse")
            m3.metric("System Adoption (App Users)", df_final[df_final['Source'] == 'App']['Name'].count())
            
            st.divider()
            c1, c2 = st.columns(2)
            with c1:
                fig_perf = px.histogram(df_final[df_final['Source'] != '-'], x="Source", color="Status", 
                                       barmode="group", title="Lateness: App vs Physical Gates",
                                       color_discrete_map={"🔴 Late": "#d73a49", "✅ On Time": "#22863a"})
                st.plotly_chart(fig_perf, use_container_width=True)
            with c2:
                fig_dist = px.pie(df_final, names='Status', title="Overall Status Distribution",
                                 color_discrete_map={"🔴 Late": "#d73a49", "✅ On Time": "#22863a", "❌ Absence": "#7a7a7a", "🟡 Weekly Off": "#ffd700"})
                st.plotly_chart(fig_dist, use_container_width=True)

    with tab2:
        st.header("Consolidated Audit Report")
        search = st.text_input("🔍 Search 500+ Employees...")
        if search:
            df_final = df_final[df_final['Name'].str.contains(search, na=False)]

        def style_rows(val):
            if "Late" in str(val) or "Absence" in str(val): return 'background-color: #ffeef0; color: #d73a49; font-weight: bold;'
            if "On Time" in str(val): return 'background-color: #e6ffed; color: #22863a; font-weight: bold;'
            if "Off" in str(val): return 'background-color: #fffbdd; color: #735c0f;'
            return ''

        st.dataframe(df_final.style.applymap(style_rows, subset=['Status']), use_container_width=True, hide_index=True)

        buf = BytesIO()
        with pd.ExcelWriter(buf, engine='xlsxwriter') as writer:
            df_final.to_excel(writer, index=False, sheet_name='Audit')
        st.download_button("📥 Export Analysis Report (Excel)", buf.getvalue(), f"HR_Report_{target_date}.xlsx")

else:
    st.info("Please upload logs to view the Analysis and Charts.")

    run_exceptions_module()
